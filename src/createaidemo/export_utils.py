"""导出工具：将 Markdown 报告转换为 Word / Markdown 文件下载"""
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor


# 行内格式 token：加粗 / 行内代码 / 链接 / 斜体
_INLINE_RE = re.compile(
    r"\*\*(.+?)\*\*"                # **bold**
    r"|`([^`]+?)`"                 # `code`
    r"|\[([^\]]+)\]\(([^)]+)\)"    # [text](url)
    r"|\*(.+?)\*"                  # *italic*
)

# 表格分隔行（| --- | --- | 或 --- | ---）
_SEP_CELL_RE = re.compile(r"^:?-+:?$")


def export_to_word(markdown_text: str, title: str = "报告") -> bytes:
    """将 Markdown 内容转换为格式化的 Word 文档（.docx）。

    支持：标题 H1-H4、加粗、斜体、行内代码、链接、有序/无序列表、
    引用、分隔线、代码块、普通段落，以及 Markdown 表格。
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块（围栏 ```）单独整块处理
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束围栏
            for cl in code_lines:
                p = doc.add_paragraph(cl)
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 表格：当前行含 | 且下一行是分隔行
        if "|" in stripped and i + 1 < n and _is_table_sep(lines[i + 1].strip()):
            header = _split_row(stripped)
            j = i + 2
            rows = []
            while j < n and "|" in lines[j] and not _is_table_sep(lines[j].strip()):
                rows.append(_split_row(lines[j].strip()))
                j += 1
            _add_table(doc, header, rows)
            doc.add_paragraph()  # 表格后留白
            i = j
            continue

        # 分隔线
        if stripped.startswith("---") or stripped.startswith("***"):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(m.group(2), level=level)
            _set_heading_font(p, Pt(20 - 2 * level))
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(16)
            _add_inline(p, stripped.lstrip("> ").strip())
            for r in p.runs:
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:].strip())
            i += 1
            continue

        # 有序列表
        om = re.match(r"^\d+\.\s+(.*)$", stripped)
        if om:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, om.group(1).strip())
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _is_table_sep(line: str) -> bool:
    """判断一行是否为 Markdown 表格分隔行（如 | --- | --- |）。"""
    s = line.strip().strip("|").strip()
    if "-" not in s:
        return False
    cells = [c.strip() for c in s.split("|")]
    return all(_SEP_CELL_RE.match(c) for c in cells if c != "")


def _split_row(line: str):
    """将表格行拆为单元格文本列表，兼容首/尾是否有管道符。"""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, header, rows):
    n_cols = max([len(header)] + [len(r) for r in rows] + [1])
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for c in range(n_cols):
        text = header[c] if c < len(header) else ""
        hdr_cells[c].text = ""
        _add_inline(hdr_cells[c].paragraphs[0], text)
        for r in hdr_cells[c].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for c in range(n_cols):
            text = row[c] if c < len(row) else ""
            cells[c].text = ""
            _add_inline(cells[c].paragraphs[0], text)


def _set_heading_font(paragraph, size: Pt):
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = size
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _add_inline(paragraph, text: str):
    """解析行内格式（加粗/斜体/代码/链接）并写入段落。"""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(1) is not None:
            r = paragraph.add_run(m.group(1))
            r.bold = True
        elif m.group(2) is not None:
            r = paragraph.add_run(m.group(2))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif m.group(3) is not None:
            paragraph.add_run(m.group(3))
            url = m.group(4)
            if url and not url.lower().startswith(("javascript:", "data:")):
                note = paragraph.add_run(f" ({url})")
                note.font.size = Pt(9)
                note.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif m.group(5) is not None:
            r = paragraph.add_run(m.group(5))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
